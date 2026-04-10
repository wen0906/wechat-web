#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件解析模块

支持解析：
- .docx 文件（Microsoft Word）
- .md 文件（Markdown）
- .txt 文件（纯文本）
"""

import os
from pathlib import Path
from typing import Tuple


class FileParser:
    """文件解析器"""
    
    SUPPORTED_EXTENSIONS = ['.docx', '.md', '.txt']
    
    @classmethod
    def parse(cls, file_path: str) -> Tuple[str, str]:
        """
        解析文件并返回 (标题, 内容)
        
        参数：
            file_path: 文件路径
            
        返回：
            (标题, 内容) 元组
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"文件不存在：{file_path}")
        
        ext = path.suffix.lower()
        
        if ext not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式：{ext}，支持的格式：{cls.SUPPORTED_EXTENSIONS}")
        
        if ext == '.docx':
            return cls._parse_docx(path)
        elif ext == '.md':
            return cls._parse_markdown(path)
        else:
            return cls._parse_text(path)
    
    @classmethod
    def _parse_docx(cls, path: Path) -> Tuple[str, str]:
        """解析 Word 文档"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("请安装 python-docx 库：pip install python-docx")
        
        doc = Document(path)
        
        # 提取标题（第一个段落）
        title = "未命名文章"
        if doc.paragraphs:
            first_para = doc.paragraphs[0].text.strip()
            if first_para:
                title = first_para
        
        # 提取正文
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        content = '\n\n'.join(paragraphs)
        
        # 尝试提取标题级别
        content = cls._enhance_headings(doc, content)
        
        return title, content
    
    @classmethod
    def _enhance_headings(cls, doc, content: str) -> str:
        """为 Word 文档添加 Markdown 标题标记"""
        lines = content.split('\n\n')
        enhanced_lines = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 根据样式名称判断标题级别
            style_name = para.style.name.lower() if para.style else ""
            
            if 'heading 1' in style_name or '标题 1' in style_name:
                enhanced_lines.append(f"# {text}")
            elif 'heading 2' in style_name or '标题 2' in style_name:
                enhanced_lines.append(f"## {text}")
            elif 'heading 3' in style_name or '标题 3' in style_name:
                enhanced_lines.append(f"### {text}")
            else:
                enhanced_lines.append(text)
        
        return '\n\n'.join(enhanced_lines)
    
    @classmethod
    def _parse_markdown(cls, path: Path) -> Tuple[str, str]:
        """解析 Markdown 文件"""
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 提取标题
        title = cls._extract_title(content)
        
        return title, content
    
    @classmethod
    def _parse_text(cls, path: Path) -> Tuple[str, str]:
        """解析纯文本文件"""
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 提取标题（第一行）
        lines = content.split('\n')
        title = lines[0].strip() if lines else "未命名文章"
        
        return title, content
    
    @classmethod
    def _extract_title(cls, content: str) -> str:
        """从 Markdown 内容中提取标题"""
        lines = content.split('\n')
        for line in lines:
            if line.strip().startswith('# '):
                return line.strip()[2:].strip()
        return "未命名文章"
    
    @classmethod
    def validate_file(cls, file_path: str) -> Tuple[bool, str]:
        """
        验证文件是否有效
        
        返回：
            (是否有效, 错误消息)
        """
        if not file_path:
            return False, "请选择文件"
        
        path = Path(file_path)
        
        if not path.exists():
            return False, f"文件不存在：{file_path}"
        
        if not path.is_file():
            return False, "请选择文件而非文件夹"
        
        ext = path.suffix.lower()
        if ext not in cls.SUPPORTED_EXTENSIONS:
            return False, f"不支持的格式：{ext}，支持：{', '.join(cls.SUPPORTED_EXTENSIONS)}"
        
        return True, ""
